"""
Resume Builder Service - PDF and DOCX Generation
"""

import io
from typing import Dict, Any, Optional
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

# DOCX imports - optional, fallback to PDF only if not available
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ResumeBuilder:
    """Service for generating ATS-friendly PDF and DOCX resumes"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom paragraph styles"""
        self.styles.add(ParagraphStyle(
            name='Name',
            fontSize=18,
            fontName='Helvetica-Bold',
            spaceAfter=6,
            textColor=colors.HexColor('#1a1a1a')
        ))
        
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            fontSize=12,
            fontName='Helvetica-Bold',
            spaceBefore=12,
            spaceAfter=6,
            textColor=colors.HexColor('#2563eb')
        ))
        
        self.styles.add(ParagraphStyle(
            name='JobTitle',
            fontSize=11,
            fontName='Helvetica-Bold',
            spaceBefore=6,
            spaceAfter=2
        ))
        
        self.styles.add(ParagraphStyle(
            name='Company',
            fontSize=10,
            fontName='Helvetica-Oblique',
            textColor=colors.HexColor('#4b5563')
        ))
        
        self.styles.add(ParagraphStyle(
            name='BulletPoint',
            fontSize=10,
            fontName='Helvetica',
            leftIndent=20,
            spaceBefore=2
        ))
    
    def _hex_to_rgb(self, hex_color: str) -> tuple:
        """Convert hex color to RGB tuple"""
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    def _format_date(self, date_str: str) -> str:
        """Format date string for display"""
        if not date_str:
            return ''
        try:
            from datetime import datetime
            date = datetime.strptime(date_str, '%Y-%m')
            return date.strftime('%b %Y')
        except:
            return date_str
    
    def generate_pdf(self, resume_content: Dict[str, Any], theme_color: str = '#3B82F6') -> bytes:
        """Generate PDF from resume content (new format)"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.5*inch,
            leftMargin=0.5*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )
        
        # Update section header color based on theme
        self.styles['SectionHeader'].textColor = colors.HexColor(theme_color)
        
        story = []
        
        # Handle both old and new resume formats
        personal = resume_content.get('personalInfo', resume_content.get('personal', {}))
        
        # Personal Info - Header
        name = personal.get('fullName') or f"{personal.get('first_name', '')} {personal.get('last_name', '')}".strip()
        if name:
            story.append(Paragraph(name, self.styles['Name']))
        
        # Contact line
        contact_parts = []
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        if personal.get('location'):
            contact_parts.append(personal['location'])
        if personal.get('linkedin') or personal.get('linkedin_url'):
            contact_parts.append(personal.get('linkedin') or 'LinkedIn')
        
        if contact_parts:
            story.append(Paragraph(' | '.join(contact_parts), self.styles['Normal']))
        
        story.append(Spacer(1, 12))
        
        # Summary
        summary = personal.get('summary') or resume_content.get('summary')
        if summary:
            story.append(Paragraph('PROFESSIONAL SUMMARY', self.styles['SectionHeader']))
            story.append(Paragraph(summary, self.styles['Normal']))
        
        # Experience
        experiences = resume_content.get('experiences', resume_content.get('experience', []))
        if experiences:
            story.append(Paragraph('EXPERIENCE', self.styles['SectionHeader']))
            for exp in experiences[:5]:
                # Handle both old and new formats
                position = exp.get('position') or exp.get('title', '')
                company = exp.get('company', '')
                
                title_company = f"<b>{position}</b> at {company}"
                story.append(Paragraph(title_company, self.styles['JobTitle']))
                
                # Date range
                start = self._format_date(exp.get('startDate', '')) or exp.get('start_date', '')
                end = 'Present' if exp.get('current') else (self._format_date(exp.get('endDate', '')) or exp.get('end_date', ''))
                duration = f"{start} - {end}" if start else exp.get('duration', '')
                if duration:
                    story.append(Paragraph(duration, self.styles['Company']))
                
                # Description/Achievements
                description = exp.get('description', '')
                if description:
                    story.append(Paragraph(f"• {description}", self.styles['BulletPoint']))
                
                achievements = exp.get('highlights', exp.get('achievements', []))
                for achievement in achievements[:3]:
                    story.append(Paragraph(f"• {achievement}", self.styles['BulletPoint']))
        
        # Education
        education = resume_content.get('education', [])
        if education:
            story.append(Paragraph('EDUCATION', self.styles['SectionHeader']))
            for edu in education[:3]:
                institution = edu.get('institution', '')
                degree = edu.get('degree', '')
                field = edu.get('field', '')
                
                edu_text = f"<b>{degree}</b>"
                if field:
                    edu_text += f" in {field}"
                if institution:
                    edu_text += f" - {institution}"
                
                story.append(Paragraph(edu_text, self.styles['Normal']))
                
                # Date range
                start = self._format_date(edu.get('startDate', ''))
                end = self._format_date(edu.get('endDate', ''))
                if start or end:
                    story.append(Paragraph(f"{start} - {end}", self.styles['Company']))
        
        # Skills
        skills = resume_content.get('skills', [])
        if skills:
            story.append(Paragraph('SKILLS', self.styles['SectionHeader']))
            skills_text = ' • '.join(skills[:20])
            story.append(Paragraph(skills_text, self.styles['Normal']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.read()
    
    def generate_docx(self, resume_content: Dict[str, Any], theme_color: str = '#3B82F6') -> bytes:
        """Generate DOCX from resume content"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not installed. Cannot generate DOCX.")
        
        document = Document()
        
        # Set narrow margins
        sections = document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Theme color
        rgb = self._hex_to_rgb(theme_color)
        theme_rgb = RGBColor(rgb[0], rgb[1], rgb[2])
        
        # Personal Info
        personal = resume_content.get('personalInfo', resume_content.get('personal', {}))
        
        # Name
        name = personal.get('fullName') or f"{personal.get('first_name', '')} {personal.get('last_name', '')}".strip()
        if name:
            name_para = document.add_heading(name, level=0)
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in name_para.runs:
                run.font.size = Pt(24)
                run.font.color.rgb = RGBColor(26, 26, 26)
        
        # Contact info
        contact_parts = []
        if personal.get('email'):
            contact_parts.append(personal['email'])
        if personal.get('phone'):
            contact_parts.append(personal['phone'])
        if personal.get('location'):
            contact_parts.append(personal['location'])
        if personal.get('linkedin'):
            contact_parts.append(personal['linkedin'])
        
        if contact_parts:
            contact_para = document.add_paragraph(' | '.join(contact_parts))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in contact_para.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Summary
        summary = personal.get('summary') or resume_content.get('summary')
        if summary:
            heading = document.add_heading('PROFESSIONAL SUMMARY', level=1)
            for run in heading.runs:
                run.font.color.rgb = theme_rgb
                run.font.size = Pt(12)
            document.add_paragraph(summary)
        
        # Experience
        experiences = resume_content.get('experiences', resume_content.get('experience', []))
        if experiences:
            heading = document.add_heading('EXPERIENCE', level=1)
            for run in heading.runs:
                run.font.color.rgb = theme_rgb
                run.font.size = Pt(12)
            
            for exp in experiences[:5]:
                position = exp.get('position') or exp.get('title', '')
                company = exp.get('company', '')
                
                # Job title line
                job_para = document.add_paragraph()
                run = job_para.add_run(f"{position}")
                run.bold = True
                run.font.size = Pt(11)
                job_para.add_run(f" at {company}")
                
                # Date range
                start = self._format_date(exp.get('startDate', '')) or exp.get('start_date', '')
                end = 'Present' if exp.get('current') else (self._format_date(exp.get('endDate', '')) or exp.get('end_date', ''))
                if start:
                    date_para = document.add_paragraph(f"{start} - {end}")
                    for run in date_para.runs:
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(100, 100, 100)
                
                # Description
                description = exp.get('description', '')
                if description:
                    bullet = document.add_paragraph(description, style='List Bullet')
                    for run in bullet.runs:
                        run.font.size = Pt(10)
                
                highlights = exp.get('highlights', exp.get('achievements', []))
                for highlight in highlights[:3]:
                    bullet = document.add_paragraph(highlight, style='List Bullet')
                    for run in bullet.runs:
                        run.font.size = Pt(10)
        
        # Education
        education = resume_content.get('education', [])
        if education:
            heading = document.add_heading('EDUCATION', level=1)
            for run in heading.runs:
                run.font.color.rgb = theme_rgb
                run.font.size = Pt(12)
            
            for edu in education[:3]:
                institution = edu.get('institution', '')
                degree = edu.get('degree', '')
                field = edu.get('field', '')
                
                edu_para = document.add_paragraph()
                run = edu_para.add_run(degree)
                run.bold = True
                if field:
                    edu_para.add_run(f" in {field}")
                if institution:
                    edu_para.add_run(f" - {institution}")
        
        # Skills
        skills = resume_content.get('skills', [])
        if skills:
            heading = document.add_heading('SKILLS', level=1)
            for run in heading.runs:
                run.font.color.rgb = theme_rgb
                run.font.size = Pt(12)
            
            skills_para = document.add_paragraph(' • '.join(skills[:20]))
            for run in skills_para.runs:
                run.font.size = Pt(10)
        
        # Save to buffer
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def analyze_ats_compatibility(self, resume_content: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze resume for ATS compatibility"""
        issues = []
        suggestions = []
        score = 100
        
        personal = resume_content.get('personalInfo', resume_content.get('personal', {}))
        
        # Check for essential sections
        if not personal.get('summary') and not resume_content.get('summary'):
            issues.append("Missing professional summary")
            score -= 10
        
        if not resume_content.get('skills'):
            issues.append("Missing skills section")
            score -= 15
        
        experiences = resume_content.get('experiences', resume_content.get('experience', []))
        if len(experiences) == 0:
            issues.append("No work experience listed")
            score -= 20
        
        if len(resume_content.get('education', [])) == 0:
            suggestions.append("Consider adding education details")
            score -= 5
        
        # Check content quality
        for exp in experiences:
            company = exp.get('company', 'your roles')
            if not exp.get('description') and not exp.get('highlights') and not exp.get('achievements'):
                suggestions.append(f"Add achievements for {company}")
        
        return {
            "score": max(0, min(100, score)),
            "issues": issues,
            "suggestions": suggestions,
            "format_issues": []
        }


# Singleton
resume_builder = ResumeBuilder()
